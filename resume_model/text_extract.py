import fitz
import docx
import re
from typing import List, Dict, Tuple

def extract_text_and_links_pdf(pdf_path: str) -> Tuple[str, List[Dict]]:
    doc = fitz.open(pdf_path)
    full_text = ''
    links = []
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        full_text += page.get_text()
        for link in page.get_links():
            if 'uri' in link:
                links.append({'page': page_num + 1, 'uri': link['uri'], 'rect': link['from']})
    return full_text.strip(), links

def extract_text_and_links_docx(docx_path: str) -> Tuple[str, List[Dict]]:
    doc = docx.Document(docx_path)
    full_text = []
    links = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            full_text.append(text)
        for run in para.runs:
            if 'HYPERLINK' in run._element.xml:
                match = re.search(r'HYPERLINK "(.*?)"', run._element.xml)
                if match:
                    uri = match.group(1)
                    links.append({'uri': uri, 'text': text})
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                full_text.append(' | '.join(row_text))
    rels = doc.part.rels
    for rel in rels.values():
        if rel.reltype == 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink':
            links.append({'uri': rel.target_ref, 'text': ''})
    return '\n'.join(full_text).strip(), links

def extract_text_and_links(file_path: str) -> Tuple[str, List[Dict]]:
    if file_path.lower().endswith('.pdf'):
        return extract_text_and_links_pdf(file_path)
    elif file_path.lower().endswith('.docx'):
        return extract_text_and_links_docx(file_path)
    else:
        raise ValueError('Unsupported file format')
